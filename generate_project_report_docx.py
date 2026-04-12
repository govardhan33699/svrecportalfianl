"""
SVREC College ERP Portal - Project Report Generator (Word .docx format)
Generates an editable project report using python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# ──────────────────────────────────────────────
# Color Palette (Hex to RGB for Word)
# ──────────────────────────────────────────────
def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))

PRIMARY_RGB   = hex_to_rgb('1a237e')   # Deep Navy Blue
SECONDARY_RGB = hex_to_rgb('0d47a1')   # Medium Blue
ACCENT_RGB    = hex_to_rgb('1565c0')   # Bright Blue
DARK_GREY_RGB = hex_to_rgb('212121')
MID_GREY_RGB  = hex_to_rgb('616161')


# ──────────────────────────────────────────────
# Helper Functions
# ──────────────────────────────────────────────
def set_cell_background(cell, fill, color_type="sRGB-clr"):
    """Sets the background color of a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1, color=PRIMARY_RGB, size=Pt(22), bold=True):
    h = doc.add_heading(text, level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(*color)
    run.font.size = size
    run.font.bold = bold
    return h

def add_paragraph(doc, text, style='Normal', bold=False, italic=False, color=DARK_GREY_RGB, size=Pt(10.5)):
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.runs[0] if p.runs else p.add_run()
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)
    run.font.size = size
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27)
    return p

def add_section_divider(doc):
    doc.add_paragraph("_" * 80)


# ══════════════════════════════════════════════
# Document Builder
# ══════════════════════════════════════════════
def generate_docx_report(output_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ═══════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════
    doc.add_paragraph('\n' * 5)
    
    p = doc.add_paragraph("SRI VENKATESWARA RAJYA LAKSMI ENGINEERING COLLEGE")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*SECONDARY_RGB)
    
    p = doc.add_paragraph("(Autonomous) — Affiliated to JNTUK")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*MID_GREY_RGB)
    
    doc.add_paragraph('\n')
    
    p = doc.add_paragraph("SVREC ERP Portal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(*PRIMARY_RGB)
    
    p = doc.add_paragraph("Enterprise Resource Planning System")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(*SECONDARY_RGB)
    
    doc.add_paragraph('\n' * 2)
    
    p = doc.add_paragraph("Comprehensive Project Report")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(*MID_GREY_RGB)
    
    doc.add_paragraph('\n' * 3)
    
    # Metadata Table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    metadata = [
        ("Project Name",     "SVREC ERP Portal"),
        ("Institution",      "Sri Venkateswara Rajya Laksmi Engineering College"),
        ("Framework",        "Django 4.2 (Python)"),
        ("Database",         "SQLite (Dev) / PostgreSQL (Production)"),
        ("Report Date",      datetime.now().strftime("%d %B %Y")),
        ("Version",          "v3.0 — Editable Production Release"),
    ]
    for i, (label, value) in enumerate(metadata):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        set_cell_background(table.cell(i, 0), 'E3F2FD')
        
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # TABLE OF CONTENTS placeholder
    # ═══════════════════════════════════════════
    add_heading(doc, "TABLE OF CONTENTS", level=1)
    toc_items = [
        "1. Project Overview",
        "2. System Architecture",
        "3. Technology Stack",
        "4. Data Models & Database Design",
        "5. Module 1 — Admin / HOD Portal",
        "6. Module 2 — Staff Faculty Portal",
        "7. Module 3 — Student Portal",
        "8. Key Advanced Features",
        "9. Security & Access Control",
        "10. Deployment & Configuration",
        "11. URL Route Map",
        "12. Conclusion & Future Scope",
    ]
    for i, item in enumerate(toc_items, 1):
        doc.add_paragraph(f"{item}", style='List Number')
    
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # CHAPTER 1: OVERVIEW
    # ═══════════════════════════════════════════
    add_heading(doc, "1. Project Overview")
    
    add_heading(doc, "1.1 Introduction", level=2, size=Pt(14), color=SECONDARY_RGB)
    add_paragraph(doc, 
        "The SVREC ERP Portal is a full-featured, web-based Enterprise Resource Planning "
        "system purpose-built for Sri Venkateswara Rajya Laksmi Engineering College (Autonomous). "
        "The system catering to three primary user roles — HOD/Admin, Faculty/Staff, and Students — "
        "each with a dedicated portal covering the full academic lifecycle."
    )
    
    add_heading(doc, "1.2 Objectives", level=2, size=Pt(14), color=SECONDARY_RGB)
    objectives = [
        "Digitize all academic and administrative workflows.",
        "Provide role-based interfaces for Administrators, Faculty, and Students.",
        "Enable real-time attendance and results management.",
        "Support automated email notifications via a workflow engine.",
        "Provide student cloud storage and certificate management."
    ]
    for obj in objectives:
        add_bullet(doc, obj)

    # ═══════════════════════════════════════════
    # CHAPTER 2: ARCHITECTURE
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "2. System Architecture")
    
    add_heading(doc, "2.1 Architectural Pattern", level=2, size=Pt(14), color=SECONDARY_RGB)
    add_paragraph(doc, 
        "Follows the Model-View-Template (MVT) architectural pattern. Business logic is in Views, "
        "data schemas in Models (ORM), and UI in HTML templates."
    )
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Responsibility'
    for cell in hdr_cells:
        set_cell_background(cell, '1A237E')
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        
    arch_data = [
        ("Presentation", "HTML5, CSS3, Bootstrap 4, JS", "User Interface"),
        ("Application", "Django 4.2 (Python)", "Business Logic"),
        ("ORM", "Django ORM", "Database Abstraction"),
        ("Database", "SQLite / PostgreSQL", "Data Persistence"),
        ("Email", "Django SMTP", "Notifications"),
    ]
    for layer, tech, resp in arch_data:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = resp

    # ═══════════════════════════════════════════
    # CHAPTER 3: TECH STACK
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "3. Technology Stack")
    
    add_heading(doc, "3.1 Backend", level=2, size=Pt(14), color=SECONDARY_RGB)
    tech_list = [
        "Django Framework (≥ 4.2.0)",
        "Python 3.x",
        "Gunicorn WSGI Server",
        "Pandas & OpenPyXL (Excel I/O)",
        "WhiteNoise (Static file serving)"
    ]
    for tech in tech_list:
        add_bullet(doc, tech)
        
    add_heading(doc, "3.2 Frontend", level=2, size=Pt(14), color=SECONDARY_RGB)
    fe_list = [
        "Bootstrap 4.x",
        "Chart.js (Analytics)",
        "Drawflow.js (Workflow UI)",
        "Select2 (Select UI)"
    ]
    for fe in fe_list:
        add_bullet(doc, fe)

    # ═══════════════════════════════════════════
    # CHAPTER 4: MODELS
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "4. Data Models & Database Design")
    
    add_paragraph(doc, "The ERP uses 25+ relational models. Key models include:")
    models_data = [
        ("CustomUser", "Core user entity (HOD/Staff/Student)"),
        ("Student", "Full profile: roll number, section, admission details"),
        ("Staff", "Faculty profile: qualification, designation, joining date"),
        ("Attendance", "Attendance session tracking"),
        ("StudentResult", "Marks entry per exam per subject"),
        ("Course / Subject", "Academic structural entities"),
        ("Timetable", "Weekly class scheduling"),
        ("Workflow", "Automation graph definitions"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model'
    hdr_cells[1].text = 'Purpose'
    for cell in hdr_cells:
        set_cell_background(cell, '1A237E')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    for m, p in models_data:
        row = table.add_row().cells
        row[0].text = m
        row[1].text = p

    # ═══════════════════════════════════════════
    # MODULES
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "5. Module 1 — Admin / HOD Portal")
    add_bullet(doc, "Academic Structure (Degrees, Courses, Subjects, Sessions)")
    add_bullet(doc, "Staff & Student Profile Management")
    add_bullet(doc, "Bulk Student Import via Excel")
    add_bullet(doc, "Advanced Marks Memo & Result Tabulation")
    add_bullet(doc, "Workflow Builder & Email Automation")
    add_bullet(doc, "Academic Calendar & Timetable Management")

    doc.add_page_break()
    add_heading(doc, "6. Module 2 — Staff Faculty Portal")
    add_bullet(doc, "Attendance Management (Grid view & Update)")
    add_bullet(doc, "Marks Entry (Manual & Excel Import)")
    add_bullet(doc, "Assignment Creation & Grading")
    add_bullet(doc, "Study Material Upload")
    add_bullet(doc, "Timetable and Announcement View")

    doc.add_page_break()
    add_heading(doc, "7. Module 3 — Student Portal")
    add_bullet(doc, "Personal Attendance Tracking")
    add_bullet(doc, "Exam Results & Consolidation View")
    add_bullet(doc, "Assignment Submission System")
    add_bullet(doc, "Personal Cloud File Storage")
    add_bullet(doc, "E-Certificates Achievement View")

    # ═══════════════════════════════════════════
    # SECURITY & ROUTES
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "8. Key Advanced Features")
    add_paragraph(doc, "Features include Drawflow-based workflow automation, visual charts, and rich data tables.")

    add_heading(doc, "9. Security & Access Control")
    add_bullet(doc, "Role-Based Access Control (RBAC)")
    add_bullet(doc, "Email-based Secure Login")
    add_bullet(doc, "CSRF & Session Protection")
    add_bullet(doc, "Right-click protection across views")

    add_heading(doc, "10. Deployment & Configuration")
    add_bullet(doc, "Platform: PythonAnywhere")
    add_bullet(doc, "Automation: GitHub Action/CI/CD ready")
    add_bullet(doc, "Media handling: Local file system (media/ folder)")

    # ═══════════════════════════════════════════
    # CONCLUSION
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "12. Conclusion & Future Scope")
    add_paragraph(doc, "The SVREC ERP Portal successfully digitizes all institutional workflows. Future scope includes a mobile app and AI-based performance prediction.")
    
    doc.add_paragraph('\n' * 2)
    p = doc.add_paragraph("Report generated for: Sri Venkateswara Rajya Laksmi Engineering College")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.runs[0]
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*MID_GREY_RGB)

    # Save
    doc.save(output_path)
    print(f"Editable Report saved to: {output_path}")


if __name__ == "__main__":
    output = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "SVREC_ERP_Portal_Project_Report.docx"
    )
    generate_docx_report(output)
